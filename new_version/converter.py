import re
from PIL import ImageFont

from docx import Document
from docx.shared import Pt, Cm
from collections import defaultdict

MAX_LINE_LEN = 70

OUT_FONT = 'Times New Roman'
OUT_FONT_BACKUP = ['times.ttf']
OUT_FONT_POINTS = 12

FACTOR_TABS = 1


def getTextDimensions(text, points, font_filename):
    try:
        font = ImageFont.truetype(font_filename, points)
    except OSError as e:
        for font in OUT_FONT_BACKUP:
            try:
                font = ImageFont.truetype(font, points)
            except OSError as e_backup:
                pass

    size = font.getsize(text.upper())

    return size


def open_file(filename):
    text = open(filename, encoding='utf-8').read()
    return text


def elan_data(file):
    # elan = elan.replace('&', '')
    # elan = elan.replace('<', '&lt;')
    # elan = elan.replace('>', '&gt;')
    elan = open_file(file).splitlines()
    transc = defaultdict(str)
    transl = defaultdict(str)
    gloss = defaultdict(str)
    comment = defaultdict(str)

    for line in elan:
        tokens = line.split('\t')
        if len(tokens) == 9:
            indices = (0, 2, 4, 8)
        else:
            indices = (0, 2, 3, 4)

        layer = tokens[indices[0]]
        time_start = tokens[indices[1]]
        time_finish = tokens[indices[2]]
        text = tokens[indices[3]]

        if layer == 'transcription':
            transc[(time_start, time_finish)] = text
        elif layer == 'translation':
            transl[(time_start, time_finish)] = text
        elif layer == 'gloss':
            gloss[(time_start, time_finish)] = text
        elif layer == 'comment':
            comment[(time_start, time_finish)] = text
    return transc, transl, gloss, comment


def to_word(pivot_dictionary):
    informant = input('введите код информанта ')
    date = input('введите дату ')
    expe = input('введите свой код ')
    name = f'eve_{informant}_{date}_{expe}'

    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    document.add_paragraph()

    table = document.add_table(rows=5, cols=2)
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12.5)

    inf = table.rows[0].cells
    inf[0].text, inf[1].text = 'Информант', informant

    exp = table.rows[1].cells
    exp[0].text, exp[1].text = 'Экспедиционер', expe

    dat = table.rows[2].cells
    dat[0].text, dat[1].text = 'Дата', date

    els = table.rows[3].cells
    els[0].text, els[1].text = 'Кто ещё был на паре', input('Кто ещё был на паре? ')

    inf = table.rows[4].cells
    inf[0].text, inf[1].text = 'Примерная тематика', input('Примерная тематика ')

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Cm(0)
    document.add_paragraph().paragraph_format.space_after = Cm(0)

    for counter, (key, value) in enumerate(pivot_dictionary.items(), start=1):
        header = f'{informant}_{date}@{expe}_{counter}'
        transcription = value[0]
        translation = value[1]
        gloss = value[2]
        comment = value[3]

        p = document.add_paragraph(header, style='List Number')
        p.paragraph_format.space_after = Cm(0.1)

        transcriptions = []
        glosses = []

        transcription_tokens = transcription.split(' ')
        glosses_tokens = gloss.split(' ')
        len_transc = len(transcription_tokens)
        len_gloss = len(glosses_tokens)
        if len_transc > len_gloss:
            glosses_tokens.extend([''] * (len_transc - len_gloss))
        elif len_gloss > len_transc:
            transcription_tokens.extend([''] * (len_gloss - len_transc))
        gl_cur_len, gl_cur_run = 0, []
        transcr_cur_len, transcr_cur_run = 0, []
        last_par_index = 0

        # accumulate transcription / glosses, until adding next glosses exceeds space
        # then begin new lines and go on
        for i, (transcription_token, gloss_token) in enumerate(
                zip(transcription_tokens, glosses_tokens)):
            if (gl_cur_len + len(gloss_token) <= MAX_LINE_LEN
                and transcr_cur_len + len(transcription_token) <= MAX_LINE_LEN):
                transcr_cur_run.append(transcription_token)
                gl_cur_run.append(gloss_token)
                transcr_cur_len += len(transcription_token)
                gl_cur_len += len(gloss_token)
            else:
                transcriptions.append(transcr_cur_run)
                glosses.append(gl_cur_run)
                last_par_index += 1

                transcr_cur_run = [transcription_token]
                gl_cur_run = [gloss_token]
                transcr_cur_len, gl_cur_len = len(transcription_token), len(gloss_token)
        else:
            if len(glosses) - 1 == last_par_index - 1:
                # if num of added lines is 1 less than needed, add remaining
                transcriptions.append(transcr_cur_run)
                glosses.append(gl_cur_run)

        # tab stops determined on the go using native length rendering with font
        for i, (transcription_line, gloss_line) in enumerate(
                zip(transcriptions, glosses)):
            print(transcription_line, gloss_line)

            left_indent = 0.5
            tab_stops = [left_indent]
            for i, (transcr, gloss) in enumerate(
                    zip(transcription_line, gloss_line), start=1):
                transcr_dim = getTextDimensions(transcr, OUT_FONT_POINTS, OUT_FONT)
                gloss_dim = getTextDimensions(gloss, OUT_FONT_POINTS, OUT_FONT)
                max_dim = max((transcr_dim[0], gloss_dim[0]))
                add_cm = (
                        FACTOR_TABS * ((max_dim // 30) * 1 + int(((max_dim % 30) / 30) * 4) / 4)
                        + 0.15
                )
                # TODO: this may interfere with line estimations from earlier
                # print(transcr, gloss, transcr_dim, gloss_dim, add_cm)
                # if add_cm < 1:
                #     add_cm = 1

                tab_stops.insert(i, tab_stops[i-1] + add_cm)

            p_transcription = document.add_paragraph()
            paragraph_format = p_transcription.paragraph_format
            paragraph_format.space_after = Cm(0)
            paragraph_format.left_indent = Cm(left_indent)
            p_transcription.add_run('\t'.join(transcription_line)).font.italic = True

            p_glosses = document.add_paragraph()
            paragraph_format = p_glosses.paragraph_format
            paragraph_format.space_after = Cm(0)
            paragraph_format.left_indent = Cm(0.5)

            for paragraph in (p_transcription, p_glosses):  # add all tab stops
                for tab_stop in tab_stops[1:]:
                    paragraph.paragraph_format.tab_stops.add_tab_stop(
                        Cm(tab_stop)
                    )

            for part in glossing('\t'.join(gloss_line)):
                if re.match(r'[a-z+]', part):
                    p_glosses.add_run(part).font.small_caps = True
                else:
                    p_glosses.add_run(part)

        p = document.add_paragraph()
        paragraph_format = p.paragraph_format
        paragraph_format.space_after = Cm(0.1)
        paragraph_format.left_indent = Cm(0.5)
        p.add_run(f'\'{translation}\'')
        p = document.add_paragraph()
        p.add_run(f'{key[0]} — {key[1]} {comment}')

    for paragraph in document.paragraphs:
        f = paragraph.style.font
        f.name = 'Times New Roman'
        f.size = Pt(12)
    document.save(f'{name}.docx')


def mapping(transc, transl, gloss, comment):
    pivot_dic = {}
    for key, value in transc.items():
        trnsl = transl[key]
        gls = gloss[key]
        cmnt = comment[key]
        pivot_dic[key] = [value, trnsl, gls, cmnt]
    return pivot_dic


def glossing(text):
    # text = text.replace(' ', '\t')
    glossed_text = re.split(r'([a-z+])', text)
    return glossed_text


def main():
    file = input('введите название илановского файла или назовите его 1.txt и нажмите Enter')
    if file == '':
        file = '1.txt'
    transc, transl, gloss, comment = elan_data(file)
    mapped_dic = mapping(transc, transl, gloss, comment)
    to_word(mapped_dic)


if __name__ == '__main__':
    main()
